#!/usr/bin/env python3
"""Generate the plain-language explainer as a .docx on the Desktop."""
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

d = docx.Document()
for s in d.sections:
    s.top_margin = s.bottom_margin = Inches(0.9)
    s.left_margin = s.right_margin = Inches(1.0)

st = d.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11)
st.paragraph_format.space_after = Pt(8)
st.paragraph_format.line_spacing = 1.15

GREY = RGBColor(0x44, 0x44, 0x44)
ACC = RGBColor(0x1F, 0x4E, 0x79)


def H(text, level=1):
    p = d.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACC
    return p


def P(text="", bold=False, italic=False, size=None, color=None, space=None):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if space is not None:
        p.paragraph_format.space_after = Pt(space)
    return p


def B(text, level=0):
    p = d.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + 0.25 * level)
    p.paragraph_format.space_after = Pt(4)
    return p


def N(text):
    p = d.add_paragraph(text, style="List Number")
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    return p


def MONO(text):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    return p


def CALLOUT(label, text):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(f"{label}  ")
    r.bold = True
    r.font.color.rgb = ACC
    p.add_run(text)
    return p


def TABLE(headers, rows, widths=None):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(9.5)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    d.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ─────────────────────────────── TITLE ───────────────────────────────
t = d.add_heading("How the BoardDocs Archive Actually Works", level=0)
for r in t.runs:
    r.font.color.rgb = ACC
P("Turning 3,000 school board PDFs into something you can ask questions of — "
  "and what went wrong along the way", italic=True, size=12, color=GREY)
P("Troy School District BoardDocs project · July 2026", size=9.5, color=GREY)

P()
P("This is written for someone who has never thought about document pipelines. "
  "There is no code in it. Every technical idea is explained with a plain analogy first, "
  "then the specifics of what we built, then what broke.")

# ─────────────────────────────── THE GOAL ───────────────────────────────
H("1. What we are actually trying to do", 1)

P("Troy School District posts its board materials to a system called BoardDocs — "
  "agendas, budget resolutions, financial statements, contracts, audit reports, ballot language. "
  "Thousands of documents going back years. It is all technically public, and it is all "
  "effectively unfindable. BoardDocs lets you browse one meeting at a time. It does not let you ask "
  "\"when did the district first disclose a structural deficit?\" or \"what has the Headlee rollback "
  "actually cost us?\"")

P("The goal is a searchable archive where those questions have answers, and where an AI assistant "
  "can read across the whole corpus rather than one document at a time.")

P("That turns out to require four separate things to work, in order:")

TABLE(
    ["Step", "What it means", "The plain version"],
    [["Crawl", "Download every document from BoardDocs",
      "Get the PDFs onto a hard drive"],
     ["Extract", "Turn each PDF into plain text",
      "Read the words out of the picture"],
     ["Chunk", "Cut the text into searchable pieces",
      "Tear the book into index cards"],
     ["Summarize", "Write a description of each document",
      "Write the blurb on the back cover"]],
    widths=[0.9, 2.6, 2.5])

P("Each step sounds trivial. Each one has a failure mode that is invisible until you go looking for it. "
  "Three of the four bit us.")

# ─────────────────────────────── EXTRACTION ───────────────────────────────
d.add_page_break()
H("2. Getting text out of a PDF is much harder than it sounds", 1)

P("Here is the thing almost nobody knows about PDFs: a PDF does not contain a document. "
  "It contains drawing instructions.")

CALLOUT("The analogy",
        "A Word file says \"here is a paragraph, then a table.\" A PDF says \"put the character 'B' "
        "at coordinates (72, 400). Put 'u' at (79, 400). Put '9' at (315, 380).\" It is closer to a "
        "painting than to a text file. Nothing in the file says which characters form a sentence, "
        "which form a table, or what order a human would read them in.")

P("So \"extracting the text\" means reconstructing a document from a picture of one. Software has to "
  "guess where words end, where columns are, and — critically — what order things go in. "
  "Different tools guess differently.")

H("The two tools, and why each one is wrong", 2)

P("We use two open-source libraries. They fail in exactly opposite ways.")

TABLE(
    ["", "pypdf", "pdfplumber"],
    [["How it reads",
      "In the order the PDF file happens to store the drawing instructions",
      "By looking at where things sit on the page, top to bottom"],
     ["Gets numbers right?",
      "Yes — perfectly",
      "No — splits digits apart"],
     ["Gets order right?",
      "No — scrambles it",
      "Yes"],
     ["Typical damage",
      "A table's heading appears AFTER the table it belongs to",
      "$123,879,792 comes out as \"1 23,879,792\""]],
    widths=[1.3, 2.4, 2.4])

P("Both failures are catastrophic for this particular archive, for different reasons.")

P("The pdfplumber failure is obvious once you see it: a budget number of $123,879,792 silently becomes "
  "two numbers, \"1\" and \"23,879,792\". In an archive whose entire purpose is answering questions about "
  "dollar amounts, that is fatal.")

P("The pypdf failure is sneakier and much more dangerous, because the text looks perfect. Consider a "
  "budget document with sections for the 2019 Capital Projects Fund and the 2023 Capital Projects Fund. "
  "pypdf might emit:")

MONO("  [2019 fund's numbers]\n"
     "  [2023 fund's numbers]\n"
     "  \"2019 CAPITAL PROJECTS FUND\"\n"
     "  \"2023 CAPITAL PROJECTS FUND\"")

P("Every number is correct. Every heading is correct. But anything reading top to bottom — a human, a "
  "search index, or an AI — will attach the 2019 numbers to the 2023 heading. You get a confident, "
  "fluent, completely wrong answer. We actually shipped this error into fourteen summaries before "
  "catching it.")

H("The fix: use each tool for the thing it is good at", 2)

CALLOUT("The analogy",
        "Imagine two people describing a page. One reads the words perfectly but jumbles the order. "
        "The other reads in the right order but garbles every number. So you take the first person's "
        "words and the second person's ordering, and match them up line by line.")

P("That is literally what the extractor now does. It runs both tools on the same page, matches each "
  "line between them by ignoring spaces, and then emits pypdf's characters in pdfplumber's order. "
  "Numbers stay exact; reading order gets fixed. If the two tools disagree too much to match "
  "confidently — below 60% of lines aligning — it falls back to pypdf alone rather than guessing.")

P("This costs real time, so it is skipped where it cannot help: whole-meeting packet files (hundreds of "
  "pages of mixed content with no consistent table structure) and anything over 15 MB. Those thresholds "
  "were not guessed — they were set by measuring every budget book and audit report in the corpus to "
  "make sure none of them fell outside.")

CALLOUT("Lesson",
        "When two tools disagree, the interesting question is not \"which one is right\" but "
        "\"what is each one right about?\" Neither library was usable alone. Combined, they are "
        "better than either vendor's product.")

# ─────────────────────────────── CHUNKING ───────────────────────────────
d.add_page_break()
H("3. Chunking: why we tear the documents up", 1)

P("Once we have clean text, we cut every document into overlapping pieces of roughly 800 words, "
  "each sharing about 100 words with the piece before it. These pieces are called chunks.")

CALLOUT("The analogy",
        "Think of an index at the back of a book. It does not point at the book — it points at pages. "
        "If the index only said \"this topic is discussed somewhere in this 400-page book,\" it would "
        "be useless. Chunks are the pages.")

P("Two questions usually come up here.")

P("Why cut it up at all?", bold=True)
P("Because search needs to return a location, not just a document. If someone searches for "
  "\"hold harmless millage,\" a result that says \"it's somewhere in this 300-page budget book\" is not "
  "an answer. Chunking lets the system point at the paragraph.")

P("Why overlap the pieces?", bold=True)
P("Because a sentence that straddles a boundary would otherwise be cut in half and become findable in "
  "neither piece. The 100-word overlap means every sentence sits complete inside at least one chunk. "
  "The cost is that the corpus is stored about 12% larger than the raw text — a real thing to know "
  "when counting, because naively adding up chunk sizes overstates the corpus by exactly that much.")

P("The Troy corpus is about 2,800 documents and roughly 44,000 chunks — an average of around 16 chunks "
  "per document, though a single budget book can be several hundred.")

# ─────────────────────────────── SUMMARIES ───────────────────────────────
d.add_page_break()
H("4. Summaries: the part that does the real work", 1)

P("Search on raw text has a well-known weakness. It matches words, not meaning. A search for "
  "\"budget cuts\" will not find a document that says \"expenditure reductions\" — same idea, no shared "
  "words.")

P("So every document also gets three written summaries, at three lengths:")

TABLE(
    ["Tier", "Length", "What it is for"],
    [["Paragraph", "2–4 sentences", "The line you see in a list of search results"],
     ["Page", "~400–600 words", "Read this instead of opening the document"],
     ["Verbose", "~900–1,600 words", "Every figure, name, date and decision — this is what search indexes"]],
    widths=[1.1, 1.5, 3.4])

P("The verbose tier is the important one, and it is deliberately dense with proper nouns and dollar "
  "amounts. Here is the trick that makes the whole archive work:")

CALLOUT("The key mechanism",
        "The summaries are fed back into the search index as if they were part of the document. "
        "So a search for \"budget cuts\" matches the summary sentence \"the district set a target of "
        "$5 million in expenditure reductions,\" and the real document surfaces — even though the "
        "document never uses the phrase \"budget cuts\". The summary acts as a translation layer "
        "between how people ask and how bureaucracies write.")

P("This is also why the summaries are written by a large AI model rather than generated mechanically. "
  "A mechanical summary (first 200 words, or most-frequent sentences) reproduces the document's own "
  "vocabulary, which defeats the entire purpose.")

# ─────────────────────────────── WHAT WENT WRONG ───────────────────────────────
d.add_page_break()
H("5. The bug that made us redo everything", 1)

P("Deep in the summarization script sat one line:")

MONO("  TEXT_CAP = 6000")

P("It meant: when handing a document to the AI to summarize, only hand over the first 6,000 characters "
  "— about two pages.")

P("For most documents this was harmless, because most board documents are short. For the documents "
  "that matter most it was ruinous. A 100-page budget book was being summarized from its cover page and "
  "table of contents. The summary would be fluent, plausible, and would describe almost none of the "
  "actual document.")

P("788 documents were affected. The worst cases were summarized from under 2% of their content — "
  "one 5,000-line intermediate school district budget was summarized from 1.58% of its text.")

CALLOUT("Why it stayed hidden",
        "Nothing failed. No error, no crash, no empty output. Every document had a summary, and every "
        "summary read perfectly well. You could only detect it by asking \"how much of the source did "
        "this summary actually see?\" — which nobody thinks to ask, because summaries are supposed to "
        "be shorter than their sources.")

P("This is the single most important lesson in the project, and it generalizes well beyond documents:")

CALLOUT("Lesson",
        "The most dangerous failures are the ones that produce confident, well-formed output. "
        "A crash tells you something is wrong. A plausible-looking wrong answer does not.")

P("The same pattern bit us twice more:")

B("A nightly automated job reported success every day for ten days while ingesting nothing at all. "
  "BoardDocs was silently blocking the datacenter it ran from; the job dutifully found zero new "
  "documents and reported a clean run. The tell was that a follow-up step which should have fired "
  "on every successful ingest had never fired once.")

B("A storage command was run with a mistyped folder path. It printed \"stored 0 summaries\" and exited "
  "reporting success — immediately after a database table had been cleared. It now refuses to exit "
  "cleanly when it stores nothing, because storing nothing is always a mistake.")

# ─────────────────────────────── VERIFICATION ───────────────────────────────
d.add_page_break()
H("6. How we check the summaries are true", 1)

P("An AI writing a summary can produce a figure that never appeared in the document. Not by lying — "
  "usually by doing helpful arithmetic, or by pattern-matching to what a similar document said.")

P("So every summary is machine-checked. The checker pulls out every number of four or more digits the "
  "summary asserts, and looks for it in the source text. In the most recent run that was 6,151 figures "
  "checked automatically.")

P("Getting this checker right took two attempts, and its failures are instructive.")

H("Problem one: the source text is messier than you expect", 2)

P("Numbers do not always appear cleanly in extracted text. Two artifacts show up constantly:")

B("Split numbers — $886,000 extracted as \"886, 000\" with a space in the middle")
B("Glued numbers — two separate figures 12,050 and 12,022 extracted as \"12,05012,022\"")

P("My first checker stripped all spaces from the source before looking. That fixed split numbers but "
  "created a new problem: it also glued together numbers that were genuinely separate, destroying the "
  "boundaries between them. Real figures got reported as missing.")

P("The result was 52 flagged figures, of which 46 were the checker's fault and 6 were real. The fix "
  "was to search for the number as a substring rather than as a whole token. Same 52 figures, now "
  "correctly sorted: 46 clean, 6 genuinely absent.")

CALLOUT("Lesson",
        "A verification tool is also software, and can also be wrong. A checker that cries wolf gets "
        "ignored, which is worse than no checker — the six real problems were sitting in plain sight, "
        "buried in 46 false alarms.")

H("Problem two: helpful arithmetic is still fabrication", 2)

P("The six genuine failures were all the same species. Given a monthly financial statement, the AI "
  "would compute something useful — \"revenue rose $6,578,509 from last month\" — and state it as though "
  "the document said it. The arithmetic was correct. The document did not contain that number.")

P("The original instruction permitted this: \"if you compute a difference, both operands must be "
  "present.\" That is a reasonable-sounding rule and it is wrong, because a reader cannot tell a "
  "computed figure from a reported one. In an archive of public financial records, that distinction "
  "is the whole ballgame.")

P("The instruction is now an outright ban: write no number that does not literally appear in the "
  "source. Describe the relationship in words if you like — \"expenditure exceeded revenue\" — but do "
  "not put a figure on it that the document does not print.")

P("The result on re-running the exact batch that had failed:")

TABLE(
    ["", "Figures", "Verified", "Derived", "Fabricated"],
    [["Before the fix", "797", "791", "—", "6"],
     ["After the fix", "873", "873", "0", "0"]],
    widths=[1.6, 1.2, 1.2, 1.2, 1.3])

P("Note the summaries got MORE specific, not less — 873 figures against 797. Banning invented "
  "arithmetic did not make the AI cautious; it redirected it toward figures actually on the page.")

# ─────────────────────────────── TOKENS ───────────────────────────────
d.add_page_break()
H("7. Tokens: what this actually costs, and why", 1)

P("Everything an AI does is billed in tokens. A token is about four characters — roughly three-quarters "
  "of a word. This page of text is about 500 tokens. A long budget book is about 200,000.")

P("If tokens were simply \"words in, words out,\" cost would be easy to reason about. They are not. "
  "There are four kinds, and they differ in price by a factor of fifty.")

TABLE(
    ["Kind", "What it is", "Relative price"],
    [["Input", "New text you send", "1×"],
     ["Output", "What the model writes back", "5×"],
     ["Cache write", "Storing context so it can be reused", "1.25×"],
     ["Cache read", "Re-reading context already stored", "0.1×"]],
    widths=[1.3, 3.2, 1.5])

P("That bottom row is the one that explains this entire project's economics.")

H("The thing nobody expects: re-reading dominates", 2)

CALLOUT("The analogy",
        "Go back to the consultant who re-reads the whole case file before answering anything. "
        "Re-reading is cheap per page — a tenth the price of new work. But by question 300 the file "
        "is enormous, and cheap-per-page times enormous-file times 300-questions is where all the "
        "money went.")

P("Measured on this project, across a five-hour working session:")

TABLE(
    ["Token type", "Raw volume", "Share of cost"],
    [["Cache read (re-reading context)", "209.8 million", "82%"],
     ["Output (actually writing)", "0.5 million", "10%"],
     ["Cache write (storing context)", "1.4 million", "7%"],
     ["Input (new material)", "1 thousand", "~0%"]],
    widths=[2.6, 1.7, 1.4])

P("Eighty-two percent of the cost was re-reading things already said. Ten percent was producing new "
  "work. That inverts the intuition almost everyone brings to this — including mine.")

H("What that means in practice", 2)

P("The cost of any single action depends on how long the conversation has gotten, not on how hard the "
  "action is. Measured across one session, cost per action tracked context size almost perfectly:")

TABLE(
    ["Conversation size", "Cost per action"],
    [["285,000 tokens", "37,000 weighted tokens"],
     ["440,000 tokens", "50,000"],
     ["613,000 tokens", "77,000"]],
    widths=[2.2, 2.2])

P("Reading a document is therefore not a one-time cost. It is a permanent tax on every subsequent "
  "action in that conversation. Reading a 25,000-token batch early is cheap; reading it late, when the "
  "conversation is already large, is expensive — and it makes everything after it more expensive too.")

P("This produces three practical rules that are not obvious:")

B("Do expensive reading early, while context is small.")
B("Batch actions together. Two commands in one message cost one re-read; two separate messages cost two.")
B("Verbosity is nearly free; round trips are not. Writing a longer summary costs little. Taking three "
  "steps instead of one costs a lot.")

CALLOUT("Lesson",
        "The instinct is to optimise what you produce. The actual lever was how many times we went "
        "back and forth against a large context. We were optimising the 10% and ignoring the 82%.")

H("Measuring it: a detour worth admitting to", 2)

P("For most of this project there was no visibility into any of the above. The reasoning went:")

N("Assume consumption cannot be measured locally.")
N("Discover the transcripts on disk do record it, and build a tool to parse them.")
N("Find the tool can measure consumption but not the limit, and ask for a manual calibration.")
N("Discover that the limit was already being written to a file on disk every single turn, by a "
   "status-line hook set up months earlier — and that a project note already said so.")

P("Steps 1 through 3 were unnecessary. The authoritative numbers had been sitting in a known file the "
  "whole time.")

CALLOUT("Lesson",
        "Before building a measurement tool, check what is already being measured. The cost of the "
        "detour was not the tool — the tool is useful. It was several rounds of reasoning from a "
        "false premise, and one unnecessary request for information already on disk.")

H("How the approach changed, in three phases", 2)

P("Phase 1 — everything in one conversation.", bold=True)
P("Read a batch, write summaries, verify, store, repeat. Simple and reliable, with a complete "
  "verification record. But every batch made the next one more expensive. Measured at roughly 586,000 "
  "weighted tokens per document, and about 38 documents per five-hour window — at which point the "
  "clock ran out with more than half the budget unspent. The constraint was serial speed, not money.")

P("Phase 2 — hand batches to separate AI workers.", bold=True)
P("Each worker starts fresh with only its own batch, so it never pays the accumulated-context tax. "
  "Cost per document fell roughly eight-fold, and workers run simultaneously. But fresh workers pay "
  "full price to build their context — no cache discount — so they were about three times more "
  "expensive each than predicted. Launching 35 at once consumed 40% of a five-hour allowance in "
  "eleven minutes.")

P("Phase 3 — release work in waves, with a stop line.", bold=True)
P("Same workers, released about eight at a time, with the meter checked between waves and a hard "
  "refusal to release another wave past a set threshold. This is what is running now. It is not faster "
  "than phase 2 at peak, but it never approaches the ceiling, so nothing dies mid-write.")

TABLE(
    ["", "Cost per document", "What limits it", "Failure mode"],
    [["One conversation", "~586,000", "Wall clock", "None — just slow"],
     ["35 workers at once", "~70,000", "Burn rate", "Hits ceiling mid-write"],
     ["Waves of 8", "~70,000", "Deliberate pacing", "None observed"]],
    widths=[1.5, 1.5, 1.5, 1.9])

CALLOUT("The meta-lesson",
        "Each phase was a reasonable response to the constraint that was visible at the time. The "
        "errors came from acting on a constraint that had not been measured — assuming budget was "
        "scarce when time was, then assuming fresh workers inherit a discount they do not get. "
        "Both were fixed by measuring rather than reasoning.")

# ─────────────────────────────── PARALLEL ───────────────────────────────
d.add_page_break()
H("8. Doing it faster: many AI workers at once", 1)

P("Re-summarizing hundreds of long documents one at a time is slow. The obvious fix is to run many AI "
  "workers in parallel, each handling a different batch. We did that. It worked, and it also taught us "
  "something counterintuitive about cost.")

H("Why parallel workers are cheaper — but not as much cheaper as expected", 2)

CALLOUT("The analogy",
        "Imagine a consultant who bills for re-reading the entire case file every time you ask a "
        "question. Early on the file is thin and questions are cheap. By question 300 the file is a "
        "banker's box, and every question — however small — costs a re-read of the whole box.")

P("That is how an AI assistant works. Everything said so far gets re-sent with every new action. "
  "By the middle of this project, roughly 82% of all cost was re-reading accumulated context, not "
  "producing new work. Cost per document was driven by how long the conversation had gotten, "
  "not by the document.")

P("A fresh worker starts with an empty file — just its one batch. So it should be far cheaper. "
  "I predicted about 30 times cheaper. The real figure was closer to 8 times, and the reason is worth "
  "understanding:")

CALLOUT("The catch",
        "Re-reading an existing conversation is heavily discounted, because the system has it cached. "
        "A brand-new worker has nothing cached, so it pays full price to build its context the first "
        "time. Fresh workers avoid the big pile — but they pay retail for the small one.")

P("Measured across 15 completed workers: each consumed roughly 4.4% of a five-hour budget allowance. "
  "My planning estimate had been 3.1%, taken from workers that were still running — a floor mistaken "
  "for an average.")

H("The mistake that came from that", 2)

P("Trusting the low estimate, we launched 35 workers at once. They burned 40% of the five-hour "
  "allowance in eleven minutes and were on track to exhaust it entirely in another six — while the "
  "allowance would not refresh for another thirty.")

P("Hitting that ceiling mid-run would have killed workers in the middle of writing their output, "
  "producing truncated files. We stopped the run instead. Nothing was lost: completed work was on "
  "disk and the system can resume where it left off.")

P("The queue was then rebuilt around three rules:")

N("Release work in waves of about 8 workers, not all at once.")
N("Check the meter between waves, and refuse to release another wave past a set line — leaving "
  "headroom so the ceiling can never land in the middle of a write.")
N("Size the waves by worker count, not by batch count. Two of the batches were enormous budget books "
  "split into sections, so one batch could mean six workers. Counting batches made waves wildly "
  "uneven — 6 workers in one, 11 in the next.")

CALLOUT("Lesson",
        "Estimate from completed work, never from work in progress. And when you cannot be confident "
        "in an estimate, build the guard rail rather than sharpening the estimate — waves and a "
        "release line made the bad estimate survivable.")

# ─────────────────────────────── SUMMARY OF LESSONS ───────────────────────────────
d.add_page_break()
H("9. The lessons, collected", 1)

P("Stripped of the specifics, these are the ideas worth carrying to any similar project.")

P("On silent failure", bold=True)
B("The dangerous failures produce confident, well-formed, wrong output. Crashes are a gift.")
B("When an automated job reports success, verify a side effect it should have produced — not its exit "
  "status. Ten days of green builds ingested nothing.")
B("An operation that does nothing should fail loudly. \"Stored 0 records\" is never a success.")

P("On verifying work", bold=True)
B("Check against the source, not against internal consistency. A parser can be perfectly "
  "self-consistent while silently dropping half its input.")
B("The verification tool is software too, and it can be wrong. A checker with a high false-alarm rate "
  "is worse than none, because the real problems hide in the noise.")
B("Correct arithmetic is not the same as a sourced fact. Ban the derivation rather than permitting "
  "the careful version of it.")

P("On tools", bold=True)
B("When two tools disagree, ask what each is right about rather than which to choose. The best "
  "extractor here is neither library — it is one library's characters in the other's ordering.")
B("Set thresholds by measuring the actual data, not by picking a round number. Every exclusion in the "
  "extractor was checked against the largest real document that had to survive it.")

P("On estimating and pacing", bold=True)
B("Estimate from completed work. A sample of in-flight work gives you a floor and reads like an average.")
B("Prefer a guard rail to a better estimate. Waves plus a release line made a 57%-wrong estimate "
  "harmless.")
B("Know which resource is actually scarce. For a long time we thought the constraint was budget; "
  "it was serial speed. Optimising the wrong one wastes both.")

P("On working with AI on factual material", bold=True)
B("Fluency is not accuracy, and fluent output is the hardest kind to audit by eye. Automate the check.")
B("Give the model a rule it can follow mechanically (\"no number not on the page\") rather than one "
  "requiring judgement (\"only reasonable inferences\").")
B("Tell it the work will be checked. Stating that the output is machine-verified figure by figure "
  "measurably changed behaviour.")

P()
P("— end —", italic=True, color=GREY)

out = "/Users/Alex/Desktop/BoardDocs_How_It_Works.docx"
d.save(out)
print(f"wrote {out}")
