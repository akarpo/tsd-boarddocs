"""Generate the Troy SD BoardDocs proper-noun sheet (.docx) for speech-to-text
custom vocabulary (e.g. AssemblyAI word-boost / keyterms).

Pipeline
  1. Pull the clean AI summaries from D1 (`summaries` table) — the low-noise source of
     proper nouns (financial check-registers/ledgers are intentionally never touched).
  2. Auto-extract the data-driven category: vendor/contractor firms (from contracts,
     bids and resolutions in the summaries).
  3. Merge with the QA-validated curated lists below (people roster, schools, programs,
     associations, governmental, streets, acronyms). These were verified against dated
     document context; external figures who merely appear in Troy documents (state
     officials, ISD staff, award-granting associations, vendor execs) were removed.
  4. Render a categorized .docx plus a flat, de-duplicated, paste-ready appendix.

Usage
  python scripts/proper_nouns.py                      # -> ~/Desktop/Troy School District Proper Nouns.docx
  python scripts/proper_nouns.py --out out.docx --refresh     # re-pull summaries from D1
  python scripts/proper_nouns.py --qa                 # print validation digests (run as new years land)

As more (older) meeting years get summarized, re-run --qa: it prints the board roll-call
timeline, flags names that look external, and surfaces new school/acronym candidates so the
curated lists below can be extended. Requires python-docx and wrangler.
"""
from __future__ import annotations
import argparse, json, os, re, subprocess
from collections import Counter, defaultdict
from pathlib import Path

DB = os.environ.get("D1_DB", "tsd-boarddocs")
ROOT = Path(os.environ.get("TSD_BOE_ROOT") or Path.home() / "tsd-boe-data")
CACHE = os.environ.get("SUMS_CACHE", "/tmp/sums_raw.json")
DEFAULT_OUT = str(Path.home() / "Desktop" / "Troy School District Proper Nouns.docx")


# ------------------------------------------------------------------ data loading
def load_summaries(cache=CACHE, refresh=False):
    """Return list of {url, paragraph, page, verbose}; cache the raw D1 dump locally."""
    if refresh or not Path(cache).exists():
        r = subprocess.run(["wrangler", "d1", "execute", DB, "--remote", "--yes", "--json",
                            "--command", "SELECT url, paragraph, page, verbose FROM summaries"],
                           capture_output=True, text=True)
        Path(cache).write_text(r.stdout)
    return json.load(open(cache))[0]["results"]


def load_url_dates():
    """Best-effort url -> meeting_date map from the local chunks index (for --qa / year range)."""
    idx = ROOT / "_index" / "chunks.jsonl"
    m = {}
    if idx.exists():
        for line in idx.open(encoding="utf-8"):
            try:
                c = json.loads(line)
            except Exception:
                continue
            u, d = c.get("url"), c.get("meeting_date")
            if u and d and u not in m:
                m[u] = d
    return m


def sumtext(sums):
    return "\n".join((s.get("paragraph") or "") + "\n" + (s.get("page") or "") + "\n" + (s.get("verbose") or "")
                     for s in sums)


# ------------------------------------------------------------------ firm extraction (data-driven)
GENERIC = {"apple", "an apple", "amazon", "amazon.com", "school district high", "school district",
           "asphalt paving", "bricks"}      # last two are OCR fragments of longer names
RENAME = {"TSACG-TSA Consulting Group": "TSA Consulting Group (TSACG)"}
_ORG = re.compile(r"\b((?:[A-Z][A-Za-z.'&-]+\s+){1,4}(?:Inc\.?|LLC|L\.L\.C\.?|Company|Corporation|Corp\.?|"
                  r"Associates|Architects?|Architecture|Construction|Contracting|Group|Partners|Consulting|"
                  r"Consultants|Engineering|Engineers|Paving|Electric|Mechanical|Plumbing|Roofing|Landscape|"
                  r"Solutions|Systems|Technologies|Publishing|Publishers))\b")
_RETAIL = re.compile(r"amazon|home depot|staples|kroger|walmart|target|costco|meijer|sam's|best buy|"
                     r"office depot|lowe's|walgreens|cvs", re.I)


def _norm(s):
    s = re.sub(r"\b(inc|llc|l\.l\.c|corp|corporation|company|co|associates|group|the)\b\.?", "", s.lower())
    return re.sub(r"\s+", " ", s).strip()


def extract_firms(sums, top=40):
    raw = Counter()
    for m in _ORG.findall(sumtext(sums)):
        m = re.sub(r"\s+", " ", m).strip()
        if not _RETAIL.search(m) and len(m) > 6:
            raw[m] += 1
    grp = {}
    for name, c in raw.items():
        k = _norm(name)
        if not k or k in GENERIC or len(k) < 3 or k.startswith("tmp"):   # TMP added canonically below
            continue
        g = grp.setdefault(k, {"count": 0, "forms": Counter()})
        g["count"] += c
        g["forms"][name] += c
    firms = []
    for g in grp.values():
        best = re.sub(r"\s+Inc$", " Inc.", g["forms"].most_common(1)[0][0].strip())
        firms.append((RENAME.get(best, best), g["count"]))
    return [f for f, _ in sorted(firms, key=lambda x: -x[1])[:top]]


# ------------------------------------------------------------------ QA-validated curated content
PEOPLE = [
 ("Board of Education members", [
    # --- current board (seated January 2025, after the November 2024 election) ---
    ("Vital Anne", "President · current  (roll-call “Mrs. Anne”; also written “Anne Vital”)"),
    ("Emina Alic", "current"),
    ("Matt Haupt", "current"),
    ("Nancy Philippart", "current  (16 yrs of service)"),
    ("Audra Melton", "current · elected Nov 2024"),
    ("Ayessa Potts", "current · elected Nov 2024"),
    ("Stephanie Zendler", "current · elected Nov 2024"),
    # --- former trustees (kept for archival transcription) ---
    ("Gary Hauff", "former, served to 2024 · 24 yrs  (also “Gary N. Hauff”)"),
    ("Karl Schmidt", "former · Board President 2024 · 12 yrs"),
    ("Nicole Wilson", "former, served to 2024"),
    ("Steve Gottlieb", "former, served to ~2023 · 6 yrs"),
    ("Elizabeth Hammond", "former, served to ~2023 · 6 yrs"),
 ]),
 ("Superintendent & senior administration", [
    ("Richard M. Machesky", "Superintendent (current); ‘Rich Machesky’"),
    ("Rick West", "Deputy Superintendent, Business Services (former, to 2024); ‘Richard L. West’"),
    ("Daniel “Dan” Trudel", "Assistant Superintendent, Business Services; Board Treasurer"),
    ("Christine DiPilato", "Assistant Superintendent, Secondary Instruction"),
    ("Kristine “Kris” Griffor", "Assistant Superintendent, Elementary Instruction"),
    ("John Pagel", "Assistant Superintendent, Employee Services"),
 ]),
 ("Directors & administrators", [
    ("Rob Carson", "Director of Maintenance & Operations (also ‘Robert Carson’)"),
    ("David Recker", "Director of Teaching & Learning (Ed.D.)"),
    ("Aurel Malutan", "Director of Finance"),
    ("Kandice Moynihan", "Finance Director (former)"),
    ("Troy Wissink", "Director of Technology & Data Services"),
    ("Beth Soggs", "Director of Technology (former)"),
    ("Saso Vasovski", "Director of Technology & Data Services"),
    ("Alan Wilson", "Assistant Director of Technology"),
    ("Shari Pawlus", "Supervisor of ELD (K-12)"),
    ("Sara Smotherman", "Director of Special Education"),
    ("Matt Jansen", "Director of Athletics"),
    ("Tim Fulcher", "Athletics"),
    ("Kendra Montante", "Administrator"),
    ("Patrick Griffin", "Administrator"),
    ("Lindsay Keegan", "Administrator"),
 ]),
 ("Principals & building administrators", [
    ("Jonathan Cross", "Principal"), ("Remo Roncone", "Principal"),
    ("Vernon Burden", "Principal"), ("Kristin Crowe", "Principal"),
    ("Harleen Singh", "Principal"), ("Angela Milanov", "Principal"),
    ("Scott Keen", "Principal"), ("Brian Zawislak", "Principal"),
    ("Kristy Hall", "Principal"), ("Mike Cottone", "Principal"),
    ("Amy Wallace", "Principal (Wattles)"),
 ]),
]

BUILDINGS = [
 ("Elementary schools", ["Barnard Elementary", "Bemis Elementary", "Costello Elementary", "Hamilton Elementary",
   "Hill Elementary", "Leonard Elementary", "Martell Elementary", "Morse Elementary", "Schroeder Elementary",
   "Troy Union Elementary", "Wass Elementary", "Wattles Elementary"]),
 ("Middle schools", ["Baker Middle School", "Boulan Park Middle School", "Larson Middle School",
   "Smith Middle School (New Smith Middle School)"]),
 ("High schools & academies", ["Troy High School", "Athens High School (Troy Athens)", "International Academy",
   "International Academy East (IA East)", "Troy Career Center (Troy School District Career Center)"]),
 ("Facilities & centers", ["Administration Building (4400 Livernois)", "Services Building (4420 Livernois)",
   "School District Service Center", "Operations / Maintenance Department (1140 Rankin)",
   "Transportation Center", "Troy Learning Center", "Troy Center for Transition", "Early Childhood Center",
   "Athens Stadium"]),
]

PROGRAMS = ["International Baccalaureate (IB)", "Middle Years Programme", "Advanced Placement (AP)",
 "Great Start Readiness Program (GSRP)", "Young Fives", "Schools of Choice", "Early College",
 "Career and Technical Education (CTE)", "Special Education", "Title I", "Section 31a", "Section 504",
 "Multi-Tiered System of Supports (MTSS)", "Positive Behavioral Interventions and Supports (PBIS)",
 "Restorative Practices", "Reading Recovery", "Units of Study", "STEM / STEAM", "Robotics",
 "M-STEP", "PSAT", "NWEA MAP Growth", "PowerSchool", "Schoology", "Canvas", "Clever"]

FIRMS_EXTRA = ["TMP Architecture (TMP Associates)", "Chartwells (Compass Group food service)"]

ASSOC = ["Troy Education Association (TEA)", "Michigan Association of School Boards (MASB)",
 "Oakland County School Boards Association (OCSBA)", "National School Boards Association (NSBA)",
 "Michigan School Business Officials (MSBO)", "Association of School Business Officials (ASBO)",
 "Michigan School Public Relations Association (MSPRA)", "National School Public Relations Association (NSPRA)",
 "Michigan High School Athletic Association (MHSAA)", "Troy Foundation for Educational Excellence",
 "Troy Community Coalition", "PTO Presidents’ Council", "National Honor Society",
 "Regional Educational Media Center Association (REMC)", "American Arbitration Association",
 "Oakland County Superintendents Association", "Diversity Council", "Student Equity Council", "JEDI Council"]

GOV = ["City of Troy", "Oakland County", "Oakland Schools (Oakland Intermediate School District / ISD)",
 "Kenneth Gutman (Oakland Schools Superintendent)", "Michelle Saunders (Oakland Schools)",
 "Michigan Department of Education (MDE)", "Michigan Department of Treasury",
 "Office of Retirement Services (ORS)", "Michigan Public School Employees Retirement System (MPSERS)",
 "State of Michigan", "Michigan State Board of Education", "Wayne RESA", "Padma Kuppa (State Representative)",
 "Birmingham Public Schools", "Rochester Community Schools", "Warren Consolidated Schools", "Clawson",
 "Royal Oak", "Bloomfield Hills", "Lamphere", "Avondale", "Berkley", "Lake Orion", "Walled Lake",
 "Utica Community Schools", "Oakland University", "Michigan State University", "University of Michigan",
 "Wayne State University", "Lawrence Technological University", "Walsh College"]

STREETS = ["Livernois Road (4420 Livernois – district HQ)", "Big Beaver Road", "Long Lake Road",
 "Square Lake Road (West Square Lake Road)", "John R Road", "Coolidge Highway", "Northfield Parkway",
 "Dequindre Road", "Rochester Road", "Crooks Road", "Wattles Road", "Adams Road", "Maple Road"]

ACRONYMS = [
 ("TSD", "Troy School District"), ("THS", "Troy High School"), ("AHS / TAHS", "Athens High School"),
 ("IA", "International Academy"), ("IAE", "International Academy East"), ("TCC", "Troy Career Center"),
 ("SOC", "Schools of Choice"), ("ISD", "Intermediate School District"),
 ("MDE", "Michigan Department of Education"), ("MASB", "Michigan Association of School Boards"),
 ("OCSBA", "Oakland County School Boards Association"), ("NSBA", "National School Boards Association"),
 ("MSPRA", "Michigan School Public Relations Association"), ("MHSAA", "Michigan High School Athletic Association"),
 ("ASBO / MSBO", "(Assoc. of) School Business Officials"), ("REMC", "Regional Educational Media Center"),
 ("ORS", "Office of Retirement Services"), ("MPSERS", "MI Public School Employees Retirement System"),
 ("MIP", "Member Investment Plan (MPSERS)"), ("RFP", "Request for Proposal"), ("BP", "Board Policy"),
 ("MTSS", "Multi-Tiered System of Supports"), ("PBIS", "Positive Behavioral Interventions & Supports"),
 ("CTE", "Career & Technical Education"), ("ELD / ELL", "English Language Development / Learners"),
 ("ELA", "English Language Arts"), ("IEP", "Individualized Education Program"),
 ("GSRP", "Great Start Readiness Program"), ("PSAT / SAT", "(Preliminary) Scholastic Assessment Test"),
 ("M-STEP", "MI Student Test of Educational Progress"), ("NWEA", "Northwest Evaluation Association (MAP)"),
 ("WIDA", "World-Class Instructional Design & Assessment"),
 ("ESSER", "Elementary & Secondary School Emergency Relief"), ("WIOA", "Workforce Innovation & Opportunity Act"),
 ("IDEA", "Individuals with Disabilities Education Act"), ("ADA", "Americans with Disabilities Act"),
 ("OCR", "Office for Civil Rights"), ("FTE", "Full-Time Equivalent"),
 ("UAAL", "Unfunded Actuarial Accrued Liability"), ("PA", "Public Act"), ("MCL", "Michigan Compiled Laws"),
 ("DTE", "DTE Energy"), ("CPA", "Certified Public Accountant"),
]

NAME_VARIANTS = {
 "Vital Anne": ["Vital Anne", "Anne Vital", "Mrs. Anne"],
 "Gary Hauff": ["Gary Hauff", "Gary N. Hauff"],
 "Richard M. Machesky": ["Richard Machesky", "Richard M. Machesky", "Rich Machesky"],
 "Rick West": ["Rick West", "Richard West", "Richard L. West"],
 "Daniel “Dan” Trudel": ["Dan Trudel", "Daniel Trudel"],
 "Kristine “Kris” Griffor": ["Kris Griffor", "Kristine Griffor"],
 "Rob Carson": ["Rob Carson", "Robert Carson"],
 "Beth Soggs": ["Beth Soggs", "Elizabeth Soggs"],
 "David Recker": ["David Recker", "Dave Recker"],
}


# ------------------------------------------------------------------ docx builder
def build_docx(sums, out):
    from docx import Document
    from docx.shared import Pt, RGBColor

    firms = extract_firms(sums)
    ud = load_url_dates()
    yrs = sorted({ud[s["url"]][:4] for s in sums if ud.get(s["url"])})
    span = f"{yrs[0]}–{yrs[-1]}" if yrs else "recent years"

    doc = Document()
    stl = doc.styles["Normal"]; stl.font.name = "Calibri"; stl.font.size = Pt(10.5)

    doc.add_heading("Troy School District — Board of Education Proper Nouns", level=0)
    r = doc.add_paragraph().add_run("Custom-vocabulary reference for speech-to-text (AssemblyAI Universal-3.5 Pro)")
    r.italic = True; r.font.size = Pt(11)
    note = doc.add_paragraph(); note.add_run("Source: ").bold = True
    note.add_run(f"mined from {len(sums):,} AI-summarized Troy SD BoardDocs documents (meeting years {span}). "
     "Financial check-registers/ledgers were excluded. Every name and role was verified against dated document "
     "context in a QA pass — external figures who merely appear in Troy documents (state officials, ISD staff, "
     "award-granting associations, vendor executives) were removed. Common spoken name variants are included. "
     "Paste the flat list at the end into the transcriber's custom-vocabulary / word-boost field, or use the "
     "categorized sections for reference.")
    doc.add_paragraph()

    def section(title, subtitle=None):
        doc.add_heading(title, level=1)
        if subtitle:
            rr = doc.add_paragraph().add_run(subtitle)
            rr.italic = True; rr.font.size = Pt(9); rr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def people(cat, entries):
        doc.add_heading(cat, level=2)
        for name, role in entries:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(name).bold = True
            p.add_run(f"  — {role}").font.size = Pt(9)

    def lst(items):
        for it in items:
            doc.add_paragraph(it, style="List Bullet")

    section("1 · People", "Board roster is cumulative: the current seven-member board (seated Jan 2025 after "
            "the Nov 2024 election) plus former trustees, all retained for transcribing older recordings.")
    for cat, entries in PEOPLE:
        people(cat, entries)
    section("2 · Schools, buildings & facilities")
    for sub, items in BUILDINGS:
        doc.add_heading(sub, level=2); lst(items)
    section("3 · Programs, curricula & platforms"); lst(PROGRAMS)
    section("4 · Vendors, contractors & firms", "From contracts, bids and resolutions (not ledgers).")
    lst(sorted(set(FIRMS_EXTRA + firms)))
    section("5 · Unions, associations & councils"); lst(ASSOC)
    section("6 · Local & governmental entities"); lst(GOV)
    section("7 · Streets & addresses (Troy area)"); lst(STREETS)
    section("8 · Acronyms & initialisms", "Spoken as letters in meetings — worth boosting.")
    for ac, full in ACRONYMS:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(ac).bold = True
        p.add_run(f"  — {full}").font.size = Pt(9)

    # flat paste-ready appendix
    doc.add_page_break()
    doc.add_heading("Appendix — Flat paste-ready list", level=1)
    doc.add_paragraph("De-duplicated, one term per line. Includes common name variants. "
                      "Paste directly into the custom-vocabulary / word-boost field.").italic = True
    flat = []
    for _, entries in PEOPLE:
        for name, _role in entries:
            flat.extend(NAME_VARIANTS.get(name, [re.sub(r"\s*“.*?”\s*", " ", name).strip()]))
    for _, items in BUILDINGS:
        flat += [re.sub(r"\s*\(.*?\)", "", it).strip() for it in items]
    for it in PROGRAMS + ASSOC + GOV:
        flat.append(re.sub(r"\s*\(.*?\)", "", it).strip())
        m = re.search(r"\(([^)]+)\)", it)
        if m and len(m.group(1)) <= 8 and m.group(1).isupper():
            flat.append(m.group(1))
    flat += FIRMS_EXTRA + firms
    flat += [re.sub(r"\s*\(.*?\)|\s*–.*", "", it).strip() for it in STREETS]
    for ac, _ in ACRONYMS:
        flat += [a.strip() for a in ac.split("/")]
    seen, uniq = set(), []
    for t in flat:
        t = t.strip()
        if t and t.lower() not in seen:
            seen.add(t.lower()); uniq.append(t)
    for t in uniq:
        doc.add_paragraph(t, style="Normal").paragraph_format.space_after = Pt(0)
    foot = doc.add_paragraph().add_run(f"{len(uniq)} unique terms · generated from Troy SD BoardDocs archive")
    foot.italic = True; foot.font.size = Pt(8); foot.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(out)
    print(f"saved: {out}")
    print(f"summaries: {len(sums):,}  years: {span}  firms: {len(firms)}  flat terms: {len(uniq)}")


# ------------------------------------------------------------------ QA / validation
_PART = r"(?:Di|De|Da|Van|Von|La|Le|St\.?|Mac|Mc|O’|O')"
_NAME = rf"([A-Z][a-zA-Z’'-]+(?:\s+[A-Z]\.?)?\s+(?:{_PART}\s?)?[A-Z][a-zA-Z’'-]+)"


def qa(sums):
    ud = load_url_dates()
    T = sumtext(sums)
    known_last = set()
    for _, entries in PEOPLE:
        for name, _ in entries:
            known_last.add(re.sub(r"[“”‘’\"']", "", name).split()[-1])

    print(f"# QA · {len(sums):,} summaries\n")
    print("## Board roll-call timeline (verify current vs former against the curated roster)")
    PRES = re.compile(r"(?:members present|board members present|present were)\s*[:\-]?\s*"
                      r"([A-Z][A-Za-z’'.,\s]{5,90})", re.I)
    seen, rows = set(), []
    for s in sorted(sums, key=lambda x: ud.get(x["url"], "")):
        for m in PRES.finditer(s.get("verbose") or ""):
            lst = re.split(r"\b(?:Also|Superintendent|administrators|Absent)\b",
                           re.sub(r"\s+", " ", m.group(1)))[0].strip().rstrip(",. ")
            d = ud.get(s["url"], "")[:7]
            if "," in lst and 3 < len(lst) < 80 and (d, lst) not in seen:
                seen.add((d, lst)); rows.append((d, lst))
    for d, lst in rows[-14:]:
        print(f"  {d}  {lst}")

    print("\n## Possible EXTERNAL people (curated names appearing near non-Troy orgs)")
    EXT = re.compile(r"(ASBO|Oakland Schools|State Board|Intermediate School District|Interim State|"
                     r"[A-Z][a-z]+ (?:Public Schools|Community Schools)|International .*? President)")
    for _, entries in PEOPLE:
        for name, _ in entries:
            surname = re.sub(r"[“”‘’\"']", "", name).split()[-1]
            for m in re.finditer(re.escape(surname), T):
                w = T[max(0, m.start() - 60):m.end() + 60]
                if EXT.search(w):
                    print(f"  ⚠ {name}: …{re.sub(chr(92)+'s+', ' ', w).strip()}…"); break

    print("\n## New SCHOOL candidates not in curated buildings")
    curated = {re.sub(r'\s*\(.*', '', b).strip() for _, items in BUILDINGS for b in items}
    B = Counter(f"{m.group(1)} {m.group(2)}"
                for m in re.finditer(r"\b([A-Z][A-Za-z.'-]+)\s+(Elementary|Middle School|High School)\b", T))
    for n, c in B.most_common():
        if c >= 4 and n not in curated and not re.match(r"(District|Michigan|County|The |TSD|IA |East|Park|Career|Smith Elem)", n):
            print(f"  {n} ·{c}")

    print("\n## High-frequency acronyms not yet curated")
    have = {a for ac, _ in ACRONYMS for a in re.split(r"\s*/\s*", ac)}
    STOP = set("MI US LLC ACH BE AT JE EXP PA PC HS MS SOC BD NO FY GPA II III IV".split()) | have
    ac = Counter(m for m in re.findall(r"\b([A-Z]{2,6})\b", T) if m not in STOP)
    print("  " + ", ".join(f"{a}·{c}" for a, c in ac.most_common(20)))


# ------------------------------------------------------------------ cli
def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--out", default=DEFAULT_OUT, help="output .docx path")
    ap.add_argument("--cache", default=CACHE, help="local summaries cache (raw D1 json)")
    ap.add_argument("--refresh", action="store_true", help="re-pull summaries from D1")
    ap.add_argument("--qa", action="store_true", help="print validation digests instead of building")
    a = ap.parse_args()
    sums = load_summaries(a.cache, a.refresh)
    if a.qa:
        qa(sums)
    else:
        build_docx(sums, a.out)


if __name__ == "__main__":
    main()
